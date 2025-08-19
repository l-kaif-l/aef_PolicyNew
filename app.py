
# # app.py
# import os, io, re, time
# from typing import List, Dict
# import streamlit as st
# from docx import Document
# from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type
# from openai import AzureOpenAI

# # ================== PAGE ==================
# st.set_page_config(page_title="KII Coder (Plain Text)", layout="wide")
# st.title("ğŸ“ KII Transcript â†’ Thematic Coding (Plain Text)")

# st.markdown("""
# Upload one or more **KII transcripts** (.docx or .txt).  
# The app classifies each segment into your **KII codebook** and returns **plain text** blocks:

# Theme: ...
# Subcodes: ...
# Insight: ...
# Quote: "..."


# You can read on-screen and download a TXT per file.
# """)

# # ================== AZURE CONFIG ==================
# # Add these in Streamlit Cloud â†’ App settings â†’ Secrets
# # AZURE_OPENAI_API_KEY = "..."
# # AZURE_OPENAI_ENDPOINT = "https://<your-resource>.openai.azure.com/"
# # AZURE_OPENAI_API_VERSION = "2024-10-21"
# # DEPLOYMENT = "gpt-4o"

# def init_client():
#     try:
#         return AzureOpenAI(
#             api_key=st.secrets["AZURE_OPENAI_API_KEY"],
#             api_version=st.secrets.get("AZURE_OPENAI_API_VERSION", "2024-10-21"),
#             azure_endpoint=st.secrets["AZURE_OPENAI_ENDPOINT"].strip()
#         ), st.secrets.get("DEPLOYMENT", "gpt-4o")
#     except Exception:
#         st.error("Missing/invalid Azure secrets. Set them in App settings â†’ Secrets.")
#         st.stop()

# client, DEPLOYMENT = init_client()

# # ================== KII CODEBOOK ==================
# KII_TAXONOMY: Dict[str, Dict[str, str]] = {
#     "Respondent Background (BCK)": {
#         "BCK_Role": "Role of respondent (farmer/CRP/leader).",
#         "BCK_Experience": "Years of farming / NF experience."
#     },
#     "Introduction and Spread of NF (SPR)": {
#         "SPR_Introduced_By": "Who introduced NF (CSA/govt/others).",
#         "SPR_Response": "Communityâ€™s initial reactions.",
#         "SPR_EarlyAdopters": "Presence/influence of early adopters."
#     },
#     "Adoption Patterns (ADO)": {
#         "ADO_Numbers": "Adopters initially vs now.",
#         "ADO_Types": "Types of farmers adopting (small/marginal/large).",
#         "ADO_DropoutReasons": "Reasons for dropping out."
#     },
#     "NF Practices and Inputs (PRC)": {
#         "PRC_BiologicalInputs": "Inputs like FYM/neem/panchagavya.",
#         "PRC_InputAccess": "Access/availability of inputs.",
#         "PRC_CropTypes": "Crops under NF / cropping changes."
#     },
#     "Challenges and Barriers (CHL)": {
#         "CHL_Labor": "Labor-related constraints.",
#         "CHL_Yield": "Yield-related concerns.",
#         "CHL_Market": "Market/premium issues.",
#         "CHL_Social": "Social stigma/peer pressure.",
#         "CHL_HealthFoodChange": "Diet/health change issues."
#     },
#     "Support Systems (SUP)": {
#         "SUP_CSA": "CSA training/demos/CRP support.",
#         "SUP_GovtSupport": "Govt schemes/benefits/gaps.",
#         "SUP_Suggestions": "Suggestions to improve support."
#     },
#     "Benefits of NF (BEN)": {
#         "BEN_CostReduction": "Reduced input costs.",
#         "BEN_SoilHealth": "Soil improvements.",
#         "BEN_HumanHealth": "Human health improvements.",
#         "BEN_IncomeStability": "Income stability/parity."
#     },
#     "Ecosystem and Policy (ECO)": {
#         "ECO_MarketPolicy": "Price support / exclusive markets.",
#         "ECO_EcosystemChange": "System/political leadership asks.",
#         "ECO_InstitutionalLink": "Links to govt institutions."
#     },
#     "Future Direction and Continuity (FUT)": {
#         "FUT_ExposureVisits": "Exposure/knowledge sharing.",
#         "FUT_TrainingNeed": "Further/refresher training needs.",
#         "FUT_Infrastructure": "Storage/input centres/infrastructure."
#     }
# }

# # ================== HELPERS ==================
# def read_docx(file) -> str:
#     doc = Document(file)
#     return "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())

# def read_txt(file) -> str:
#     return file.read().decode("utf-8", errors="ignore")

# def segment_text(text: str, max_chars: int = 1800) -> List[str]:
#     text = re.sub(r'\r\n?', '\n', text)
#     blocks = [b.strip() for b in re.split(r'\n{2,}', text) if b.strip()]
#     segs: List[str] = []
#     for blk in blocks:
#         if len(blk) <= max_chars:
#             segs.append(blk)
#         else:
#             # soft split on sentence boundaries
#             sentences = re.split(r'(?<=[.!?])\s+', blk)
#             cur, cur_len = [], 0
#             for s in sentences:
#                 if cur_len + len(s) > max_chars and cur:
#                     segs.append(" ".join(cur).strip())
#                     cur, cur_len = [], 0
#                 cur.append(s)
#                 cur_len += len(s) + 1
#             if cur:
#                 segs.append(" ".join(cur).strip())
#     return segs

# def taxonomy_lines(tax: Dict[str, Dict[str,str]]) -> str:
#     return "\n".join(f"{theme} â†’ {', '.join(subs.keys())}" for theme, subs in tax.items())

# def build_prompt(segment: str, tax: Dict[str,Dict[str,str]]) -> str:
#     tax_text = taxonomy_lines(tax)
#     return f"""
# You are a qualitative coding assistant. Classify the KII excerpt into the taxonomy BELOW.
# Pick exactly ONE theme and one or more subcodes from THAT theme only.
# Then write a short insight and one short verbatim quote from the text (<=25 words).

# TAXONOMY:
# {tax_text}

# TEXT:
# \"\"\"{segment}\"\"\"

# Return PLAIN TEXT only (no JSON, no backticks), exactly in this format:

# Theme: <one theme name exactly as in taxonomy>
# Subcodes: <comma-separated subcodes from the chosen theme>
# Insight: <one concise, specific sentence>
# Quote: "<short verbatim quote from TEXT, <=25 words>"

# (Do not add any extra lines or sections.)
# """.strip()

# @retry(wait=wait_exponential(multiplier=1, min=1, max=20),
#        stop=stop_after_attempt(5),
#        retry=retry_if_exception_type(Exception))
# def call_azure(prompt: str) -> str:
#     res = client.chat.completions.create(
#         model=DEPLOYMENT,
#         messages=[{"role":"user","content":prompt}],
#         temperature=0
#     )
#     return res.choices[0].message.content.strip()

# # ================== SIDEBAR ==================
# with st.sidebar:
#     uploaded_files = st.file_uploader("Upload KII transcripts (.docx / .txt)", type=["docx","txt"], accept_multiple_files=True)
#     max_chars = st.slider("Max segment size (chars)", 800, 3000, 1800, 100)
#     run = st.button("ğŸš€ Run Coding")

# # ================== MAIN ==================
# if run and uploaded_files:
#     for up in uploaded_files:
#         st.write(f"**Processing:** `{up.name}`")
#         # read file
#         if up.name.lower().endswith(".docx"):
#             text = read_docx(up)
#         else:
#             text = read_txt(up)

#         segments = segment_text(text, max_chars=max_chars)

#         # build report
#         lines = [f"=== KII CODED REPORT â€” {up.name} ===", ""]
#         progress = st.progress(0)
#         for i, seg in enumerate(segments, start=1):
#             prompt = build_prompt(seg, KII_TAXONOMY)
#             try:
#                 out = call_azure(prompt)
#             except Exception as e:
#                 out = f"Theme: (error)\nSubcodes: \nInsight: Azure error: {e}\nQuote: \"\""

#             lines.append(f"[Segment {i}]")
#             lines.append(out)
#             lines.append("")
#             progress.progress(i / max(len(segments), 1))
#             time.sleep(0.1)

#         report = "\n".join(lines)

#         # show & download
#         with st.expander(f"ğŸ“„ Report: {up.name}", expanded=True):
#             st.text(report)

#         fname_base = os.path.splitext(up.name)[0]
#         st.download_button(
#             label=f"ğŸ’¾ Download TXT â€” {fname_base}_KII_CODED.txt",
#             data=report.encode("utf-8"),
#             file_name=f"{fname_base}_KII_CODED.txt",
#             mime="text/plain"
#         )

# elif run and not uploaded_files:
#     st.warning("Please upload at least one file.")

# app.py



# ------------------------------------------------------------------------------------------------------------------------------------
# app.py
# import os, io, re, time, string
# from typing import List, Dict
# import streamlit as st
# from docx import Document
# from openai import AzureOpenAI
# from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type

# # =========================
# # Page & Intro
# # =========================
# st.set_page_config(page_title="NF Transcript Toolkit", layout="wide")
# st.title("ğŸŒ¾ NF Transcript Toolkit")
# st.caption("Switch between **Insights Extraction** and **Codebook Coding (KII)**. Azure keys are read from Streamlit Secrets.")

# # =========================
# # Azure client (from Secrets)
# # =========================
# def init_client():
#     try:
#         client = AzureOpenAI(
#             api_key=st.secrets["AZURE_OPENAI_API_KEY"],
#             api_version=st.secrets.get("AZURE_OPENAI_API_VERSION", "2024-10-21"),
#             azure_endpoint=st.secrets["AZURE_OPENAI_ENDPOINT"].strip(),  # e.g. https://<resource>.openai.azure.com/
#         )
#         deployment = st.secrets.get("DEPLOYMENT", "gpt-4o")
#         return client, deployment
#     except Exception:
#         st.error("Azure OpenAI secrets missing/invalid. Set them in App â†’ Settings â†’ Secrets.")
#         st.stop()

# client, DEPLOYMENT = init_client()

# @retry(wait=wait_exponential(multiplier=1, min=1, max=20),
#        stop=stop_after_attempt(6),
#        retry=retry_if_exception_type(Exception))
# def call_azure(prompt: str) -> str:
#     res = client.chat.completions.create(
#         model=DEPLOYMENT,
#         messages=[{"role": "user", "content": prompt}],
#         temperature=0
#     )
#     return res.choices[0].message.content.strip()

# # =========================
# # Shared helpers
# # =========================
# def read_docx_bytes(b: bytes) -> str:
#     doc = Document(io.BytesIO(b))
#     return "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())

# def read_txt_bytes(b: bytes) -> str:
#     return b.decode("utf-8", errors="ignore")

# def clean_text(s: str) -> str:
#     return ''.join(c for c in s if c in string.printable)

# def chunk_text(text: str, max_chars: int = 6000, overlap: int = 400) -> List[str]:
#     text = re.sub(r'\r\n?', '\n', text)
#     if len(text) <= max_chars:
#         return [text]
#     chunks, i = [], 0
#     while i < len(text):
#         chunk = text[i:i+max_chars]
#         lb = chunk.rfind("\n\n")
#         if lb > max_chars * 0.6:
#             chunk = chunk[:lb]
#             i += lb
#         else:
#             i += max_chars
#         chunks.append(chunk.strip())
#         i -= overlap
#         if i < 0:
#             i = 0
#     # de-dup by first 200 chars
#     seen, dedup = set(), []
#     for c in chunks:
#         k = c[:200]
#         if k not in seen:
#             seen.add(k)
#             dedup.append(c)
#     return dedup

# def normalize_bullets(lines: List[str]) -> List[str]:
#     out, seen = [], set()
#     for ln in lines:
#         s = re.sub(r'\s+', ' ', ln.strip().rstrip('.'))
#         if not s:
#             continue
#         key = s.lower()
#         if key not in seen:
#             seen.add(key)
#             out.append(s if s.endswith('.') else s + '.')
#     return out

# # =========================
# # Mode 1 â€” Insights Extraction
# # =========================
# def chunk_prompt(chunk: str) -> str:
#     return f"""
# You are a senior qualitative researcher. Read the TEXT and extract clear, non-generic insights.
# Focus on perceptions, experiences, outcomes, barriers, enablers, and actionable suggestions.
# Return ONLY the following two sections as plain text (no JSON, no Markdown):

# INSIGHTS:
# - <8â€“15 concise bullets, each a standalone insight; avoid repetition; specific to this text>

# QUOTES:
# - <5â€“10 short verbatim quotes (<=25 words each); keep any speaker tags like 'P3:' if present>

# TEXT:
# \"\"\"{chunk}\"\"\"
# """

# def synthesis_prompt(all_insights: List[str], all_quotes: List[str], fname: str) -> str:
#     ins = "\n".join(f"- {i}" for i in all_insights)
#     qts = "\n".join(f"- {q}" for q in all_quotes)
#     return f"""
# You will merge bullets from multiple chunks into a single clean report for {fname}.
# Deduplicate, group logically, and keep only the strongest items. Be specific; no filler.
# Return ONLY plain text (no JSON/Markdown), exactly in this template:

# INSIGHTS (Top 10â€“20):
# - <bullet>
# - <bullet>
# ...

# QUOTES (5â€“12 strongest):
# - <short quote>
# - <short quote>
# ...

# BULLET POOL:
# {ins}

# QUOTE POOL:
# {qts}
# """

# def run_insights(file_name: str, content: bytes) -> str:
#     text = read_docx_bytes(content) if file_name.lower().endswith(".docx") else read_txt_bytes(content)
#     text = clean_text(text)
#     chunks = chunk_text(text, max_chars=6000, overlap=400)
#     pooled_insights, pooled_quotes = [], []

#     for i, ch in enumerate(chunks, start=1):
#         resp = call_azure(chunk_prompt(ch))
#         parts = re.split(r'\bQUOTES:\s*', resp, flags=re.I)
#         insights_block = parts[0]
#         quotes_block = parts[1] if len(parts) > 1 else ""
#         ins = re.findall(r'^\s*-\s+(.*)', insights_block, flags=re.M)
#         qts = re.findall(r'^\s*-\s+(.*)', quotes_block, flags=re.M)
#         pooled_insights.extend(ins)
#         pooled_quotes.extend(qts)
#         st.progress(i/len(chunks), text=f"Extracting {i}/{len(chunks)}")

#     pooled_insights = normalize_bullets(pooled_insights)
#     pooled_quotes = normalize_bullets(pooled_quotes)
#     final_text = call_azure(synthesis_prompt(pooled_insights, pooled_quotes, file_name)).strip()
#     return f"INSIGHTS REPORT â€” {file_name}\n\n{final_text}"

# # =========================
# # Mode 2 â€” Codebook Coding (KII)
# # =========================
# KII_TAXONOMY: Dict[str, Dict[str, str]] = {
#     "Respondent Background (BCK)": {
#         "BCK_Role": "Role of respondent (farmer/CRP/leader).",
#         "BCK_Experience": "Years of farming / NF experience."
#     },
#     "Introduction and Spread of NF (SPR)": {
#         "SPR_Introduced_By": "Who introduced NF (CSA/govt/others).",
#         "SPR_Response": "Communityâ€™s initial reactions.",
#         "SPR_EarlyAdopters": "Presence/influence of early adopters."
#     },
#     "Adoption Patterns (ADO)": {
#         "ADO_Numbers": "Adopters initially vs now.",
#         "ADO_Types": "Types of farmers adopting (small/marginal/large).",
#         "ADO_DropoutReasons": "Reasons for dropping out."
#     },
#     "NF Practices and Inputs (PRC)": {
#         "PRC_BiologicalInputs": "Inputs like FYM/neem/panchagavya.",
#         "PRC_InputAccess": "Access/availability of inputs.",
#         "PRC_CropTypes": "Crops under NF / cropping changes."
#     },
#     "Challenges and Barriers (CHL)": {
#         "CHL_Labor": "Labor-related constraints.",
#         "CHL_Yield": "Yield-related concerns.",
#         "CHL_Market": "Market/premium issues.",
#         "CHL_Social": "Social stigma/peer pressure.",
#         "CHL_HealthFoodChange": "Diet/health change issues."
#     },
#     "Support Systems (SUP)": {
#         "SUP_CSA": "CSA training/demos/CRP support.",
#         "SUP_GovtSupport": "Govt schemes/benefits/gaps.",
#         "SUP_Suggestions": "Suggestions to improve support."
#     },
#     "Benefits of NF (BEN)": {
#         "BEN_CostReduction": "Reduced input costs.",
#         "BEN_SoilHealth": "Soil improvements.",
#         "BEN_HumanHealth": "Human health improvements.",
#         "BEN_IncomeStability": "Income stability/parity."
#     },
#     "Ecosystem and Policy (ECO)": {
#         "ECO_MarketPolicy": "Price support / exclusive markets.",
#         "ECO_EcosystemChange": "System/political leadership asks.",
#         "ECO_InstitutionalLink": "Links to govt institutions."
#     },
#     "Future Direction and Continuity (FUT)": {
#         "FUT_ExposureVisits": "Exposure/knowledge sharing.",
#         "FUT_TrainingNeed": "Further/refresher training needs.",
#         "FUT_Infrastructure": "Storage/input centres/infrastructure."
#     }
# }

# def taxonomy_lines(tax: Dict[str, Dict[str,str]]) -> str:
#     return "\n".join(f"{theme} â†’ {', '.join(subs.keys())}" for theme, subs in tax.items())

# def build_kII_prompt(segment: str, tax: Dict[str,Dict[str,str]]) -> str:
#     tax_text = taxonomy_lines(tax)
#     return f"""
# You are a qualitative coding assistant. Classify the KII excerpt into the taxonomy BELOW.
# Pick exactly ONE theme and one or more subcodes from THAT theme only.
# Then write a short insight and one short verbatim quote from the text (<=25 words).

# TAXONOMY:
# {tax_text}

# TEXT:
# \"\"\"{segment}\"\"\"

# Return PLAIN TEXT only (no JSON, no backticks), exactly in this format:

# Theme: <one theme name exactly as in taxonomy>
# Subcodes: <comma-separated subcodes from the chosen theme>
# Insight: <one concise, specific sentence>
# Quote: "<short verbatim quote from TEXT, <=25 words>"

# (Do not add any extra lines or sections.)
# """.strip()

# def segment_for_coding(text: str, max_chars: int = 1800) -> List[str]:
#     text = re.sub(r'\r\n?', '\n', text)
#     blocks = [b.strip() for b in re.split(r'\n{2,}', text) if b.strip()]
#     segs: List[str] = []
#     for blk in blocks:
#         if len(blk) <= max_chars:
#             segs.append(blk)
#         else:
#             sentences = re.split(r'(?<=[.!?])\s+', blk)
#             cur, cur_len = [], 0
#             for s in sentences:
#                 if cur_len + len(s) > max_chars and cur:
#                     segs.append(" ".join(cur).strip())
#                     cur, cur_len = [], 0
#                 cur.append(s)
#                 cur_len += len(s) + 1
#             if cur:
#                 segs.append(" ".join(cur).strip())
#     return segs

# def run_coding(file_name: str, content: bytes, max_chars: int) -> str:
#     text = read_docx_bytes(content) if file_name.lower().endswith(".docx") else read_txt_bytes(content)
#     segments = segment_for_coding(text, max_chars=max_chars)
#     lines = [f"=== KII CODED REPORT â€” {file_name} ===", ""]
#     for i, seg in enumerate(segments, start=1):
#         out = call_azure(build_kII_prompt(seg, KII_TAXONOMY))
#         lines.append(f"[Segment {i}]")
#         lines.append(out)
#         lines.append("")
#         st.progress(i/len(segments), text=f"Coding {i}/{len(segments)}")
#         time.sleep(0.05)
#     return "\n".join(lines)

# # =========================
# # Sidebar (Mode + Upload)
# # =========================
# st.sidebar.title("Mode")
# mode = st.sidebar.radio("Choose task:", ["Insights Extraction", "Codebook Coding (KII)"])
# uploaded_files = st.sidebar.file_uploader("Upload transcripts (.docx / .txt)", type=["docx","txt"], accept_multiple_files=True)

# # =========================
# # Main â€” Mode routing
# # =========================
# if mode == "Insights Extraction":
#     st.header("Transcript â†’ Insights Extraction")
#     st.markdown("Generates one insights report per file (bullets + short quotes).")
#     if uploaded_files and st.button("ğŸš€ Generate Insights"):
#         for up in uploaded_files:
#             st.write(f"**Processing:** `{up.name}`")
#             content = up.read()
#             report = run_insights(up.name, content)
#             with st.expander(f"ğŸ“„ Insights: {up.name}", expanded=True):
#                 st.text(report)
#             st.download_button(
#                 label=f"ğŸ’¾ Download â€” {os.path.splitext(up.name)[0]}_INSIGHTS.txt",
#                 data=report.encode("utf-8"),
#                 file_name=f"{os.path.splitext(up.name)[0]}_INSIGHTS.txt",
#                 mime="text/plain"
#             )
#     elif st.button("ğŸš€ Generate Insights") and not uploaded_files:
#         st.warning("Please upload at least one file.")

# else:
#     st.header("KII Transcript â†’ Thematic Coding (Plain Text)")
#     st.markdown("Returns Theme/Subcodes/Insight/Quote per segment.")
#     seg_size = st.slider("Max segment size (chars)", 800, 3000, 1800, 100)
#     if uploaded_files and st.button("ğŸš€ Run Coding"):
#         for up in uploaded_files:
#             st.write(f"**Processing:** `{up.name}`")
#             content = up.read()
#             report = run_coding(up.name, content, max_chars=seg_size)
#             with st.expander(f"ğŸ“„ Coded Report: {up.name}", expanded=True):
#                 st.text(report)
#             st.download_button(
#                 label=f"ğŸ’¾ Download â€” {os.path.splitext(up.name)[0]}_KII_CODED.txt",
#                 data=report.encode("utf-8"),
#                 file_name=f"{os.path.splitext(up.name)[0]}_KII_CODED.txt",
#                 mime="text/plain"
#             )
#     elif st.button("ğŸš€ Run Coding") and not uploaded_files:
#         st.warning("Please upload at least one file.")






# --------------------------------------------------------------------------------------------------------------------------------------------



# # app.py
# import os, io, re, time, string
# from typing import List, Dict
# import streamlit as st
# from docx import Document
# from openai import AzureOpenAI
# from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type

# # â”€â”€â”€â”€â”€â”€â”€â”€â”€ Page â”€â”€â”€â”€â”€â”€â”€â”€â”€
# st.set_page_config(page_title="NF Transcript Toolkit", layout="wide")
# st.title("ğŸŒ¾ NF Transcript Toolkit")
# st.caption("Use the tabs below to switch between **Insights Extraction** and **Codebook Coding (KII)**. Keys come from Streamlit Secrets.")

# # â”€â”€â”€â”€â”€â”€â”€â”€â”€ Azure Client (from Secrets) â”€â”€â”€â”€â”€â”€â”€â”€â”€
# def init_client():
#     try:
#         client = AzureOpenAI(
#             api_key=st.secrets["AZURE_OPENAI_API_KEY"],
#             api_version=st.secrets.get("AZURE_OPENAI_API_VERSION", "2024-10-21"),
#             azure_endpoint=st.secrets["AZURE_OPENAI_ENDPOINT"].strip()  # e.g. https://<resource>.openai.azure.com/
#         )
#         deployment = st.secrets.get("DEPLOYMENT", "gpt-4o")
#         return client, deployment
#     except Exception:
#         st.error("Azure OpenAI secrets missing/invalid. Set them in App â†’ Settings â†’ Secrets.")
#         st.stop()

# client, DEPLOYMENT = init_client()

# @retry(wait=wait_exponential(multiplier=1, min=1, max=20),
#        stop=stop_after_attempt(6),
#        retry=retry_if_exception_type(Exception))
# def call_azure(prompt: str) -> str:
#     res = client.chat.completions.create(
#         model=DEPLOYMENT,
#         messages=[{"role": "user", "content": prompt}],
#         temperature=0
#     )
#     return res.choices[0].message.content.strip()

# # â”€â”€â”€â”€â”€â”€â”€â”€â”€ Shared Helpers â”€â”€â”€â”€â”€â”€â”€â”€â”€
# def read_docx_bytes(b: bytes) -> str:
#     doc = Document(io.BytesIO(b))
#     return "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())

# def read_txt_bytes(b: bytes) -> str:
#     return b.decode("utf-8", errors="ignore")

# def clean_text(s: str) -> str:
#     return ''.join(c for c in s if c in string.printable)

# def chunk_text(text: str, max_chars: int = 6000, overlap: int = 400) -> List[str]:
#     text = re.sub(r'\r\n?', '\n', text)
#     if len(text) <= max_chars:
#         return [text]
#     chunks, i = [], 0
#     while i < len(text):
#         chunk = text[i:i+max_chars]
#         lb = chunk.rfind("\n\n")
#         if lb > max_chars * 0.6:
#             chunk = chunk[:lb]; i += lb
#         else:
#             i += max_chars
#         chunks.append(chunk.strip())
#         i -= overlap
#         if i < 0: i = 0
#     # de-dup by first 200 chars
#     seen, dedup = set(), []
#     for c in chunks:
#         k = c[:200]
#         if k not in seen:
#             seen.add(k); dedup.append(c)
#     return dedup

# def normalize_bullets(lines: List[str]) -> List[str]:
#     out, seen = [], set()
#     for ln in lines:
#         s = re.sub(r'\s+', ' ', ln.strip().rstrip('.'))
#         if not s: continue
#         key = s.lower()
#         if key not in seen:
#             seen.add(key)
#             out.append(s if s.endswith('.') else s + '.')
#     return out

# # â”€â”€â”€â”€â”€â”€â”€â”€â”€ Mode 1: Insights Extraction â”€â”€â”€â”€â”€â”€â”€â”€â”€
# def chunk_prompt(chunk: str) -> str:
#     return f"""
# You are a senior qualitative researcher. Read the TEXT and extract clear, non-generic insights.
# Focus on perceptions, experiences, outcomes, barriers, enablers, and actionable suggestions.
# Return ONLY the following two sections as plain text (no JSON, no Markdown):

# INSIGHTS:
# - <8â€“15 concise bullets, each a standalone insight; avoid repetition; specific to this text>

# QUOTES:
# - <5â€“10 short verbatim quotes (<=25 words each); keep any speaker tags like 'P3:' if present>

# TEXT:
# \"\"\"{chunk}\"\"\"
# """

# def synthesis_prompt(all_insights: List[str], all_quotes: List[str], fname: str) -> str:
#     ins = "\n".join(f"- {i}" for i in all_insights)
#     qts = "\n".join(f"- {q}" for q in all_quotes)
#     return f"""
# You will merge bullets from multiple chunks into a single clean report for {fname}.
# Deduplicate, group logically, and keep only the strongest items. Be specific; no filler.
# Return ONLY plain text (no JSON/Markdown), exactly in this template:

# INSIGHTS (Top 10â€“20):
# - <bullet>
# - <bullet>
# ...

# QUOTES (5â€“12 strongest):
# - <short quote>
# - <short quote>
# ...

# BULLET POOL:
# {ins}

# QUOTE POOL:
# {qts}
# """

# def run_insights(file_name: str, content: bytes) -> str:
#     text = read_docx_bytes(content) if file_name.lower().endswith(".docx") else read_txt_bytes(content)
#     text = clean_text(text)
#     chunks = chunk_text(text, max_chars=6000, overlap=400)
#     pooled_insights, pooled_quotes = [], []

#     progress = st.progress(0, text="Extracting segmentsâ€¦")
#     for i, ch in enumerate(chunks, start=1):
#         resp = call_azure(chunk_prompt(ch))
#         parts = re.split(r'\bQUOTES:\s*', resp, flags=re.I)
#         insights_block = parts[0]
#         quotes_block = parts[1] if len(parts) > 1 else ""
#         ins = re.findall(r'^\s*-\s+(.*)', insights_block, flags=re.M)
#         qts = re.findall(r'^\s*-\s+(.*)', quotes_block, flags=re.M)
#         pooled_insights.extend(ins); pooled_quotes.extend(qts)
#         progress.progress(i / len(chunks), text=f"Extracted {i}/{len(chunks)}")

#     pooled_insights = normalize_bullets(pooled_insights)
#     pooled_quotes = normalize_bullets(pooled_quotes)
#     final_text = call_azure(synthesis_prompt(pooled_insights, pooled_quotes, file_name)).strip()
#     return f"INSIGHTS REPORT â€” {file_name}\n\n{final_text}"

# # â”€â”€â”€â”€â”€â”€â”€â”€â”€ Mode 2: Codebook Coding (KII) â”€â”€â”€â”€â”€â”€â”€â”€â”€
# KII_TAXONOMY: Dict[str, Dict[str, str]] = {
#     "Respondent Background (BCK)": {
#         "BCK_Role": "Role of respondent (farmer/CRP/leader).",
#         "BCK_Experience": "Years of farming / NF experience."
#     },
#     "Introduction and Spread of NF (SPR)": {
#         "SPR_Introduced_By": "Who introduced NF (CSA/govt/others).",
#         "SPR_Response": "Communityâ€™s initial reactions.",
#         "SPR_EarlyAdopters": "Presence/influence of early adopters."
#     },
#     "Adoption Patterns (ADO)": {
#         "ADO_Numbers": "Adopters initially vs now.",
#         "ADO_Types": "Types of farmers adopting (small/marginal/large).",
#         "ADO_DropoutReasons": "Reasons for dropping out."
#     },
#     "NF Practices and Inputs (PRC)": {
#         "PRC_BiologicalInputs": "Inputs like FYM/neem/panchagavya.",
#         "PRC_InputAccess": "Access/availability of inputs.",
#         "PRC_CropTypes": "Crops under NF / cropping changes."
#     },
#     "Challenges and Barriers (CHL)": {
#         "CHL_Labor": "Labor-related constraints.",
#         "CHL_Yield": "Yield-related concerns.",
#         "CHL_Market": "Market/premium issues.",
#         "CHL_Social": "Social stigma/peer pressure.",
#         "CHL_HealthFoodChange": "Diet/health change issues."
#     },
#     "Support Systems (SUP)": {
#         "SUP_CSA": "CSA training/demos/CRP support.",
#         "SUP_GovtSupport": "Govt schemes/benefits/gaps.",
#         "SUP_Suggestions": "Suggestions to improve support."
#     },
#     "Benefits of NF (BEN)": {
#         "BEN_CostReduction": "Reduced input costs.",
#         "BEN_SoilHealth": "Soil improvements.",
#         "BEN_HumanHealth": "Human health improvements.",
#         "BEN_IncomeStability": "Income stability/parity."
#     },
#     "Ecosystem and Policy (ECO)": {
#         "ECO_MarketPolicy": "Price support / exclusive markets.",
#         "ECO_EcosystemChange": "System/political leadership asks.",
#         "ECO_InstitutionalLink": "Links to govt institutions."
#     },
#     "Future Direction and Continuity (FUT)": {
#         "FUT_ExposureVisits": "Exposure/knowledge sharing.",
#         "FUT_TrainingNeed": "Further/refresher training needs.",
#         "FUT_Infrastructure": "Storage/input centres/infrastructure."
#     }
# }

# def taxonomy_lines(tax: Dict[str, Dict[str,str]]) -> str:
#     return "\n".join(f"{theme} â†’ {', '.join(subs.keys())}" for theme, subs in tax.items())

# def build_kII_prompt(segment: str, tax: Dict[str,Dict[str,str]]) -> str:
#     tax_text = taxonomy_lines(tax)
#     return f"""
# You are a qualitative coding assistant. Classify the KII excerpt into the taxonomy BELOW.
# Pick exactly ONE theme and one or more subcodes from THAT theme only.
# Then write a short insight and one short verbatim quote from the text (<=25 words).

# TAXONOMY:
# {tax_text}

# TEXT:
# \"\"\"{segment}\"\"\"

# Return PLAIN TEXT only (no JSON, no backticks), exactly in this format:

# Theme: <one theme name exactly as in taxonomy>
# Subcodes: <comma-separated subcodes from the chosen theme>
# Insight: <one concise, specific sentence>
# Quote: "<short verbatim quote from TEXT, <=25 words>"

# (Do not add any extra lines or sections.)
# """.strip()

# def segment_for_coding(text: str, max_chars: int = 1800) -> List[str]:
#     text = re.sub(r'\r\n?', '\n', text)
#     blocks = [b.strip() for b in re.split(r'\n{2,}', text) if b.strip()]
#     segs: List[str] = []
#     for blk in blocks:
#         if len(blk) <= max_chars:
#             segs.append(blk)
#         else:
#             sentences = re.split(r'(?<=[.!?])\s+', blk)
#             cur, cur_len = [], 0
#             for s in sentences:
#                 if cur_len + len(s) > max_chars and cur:
#                     segs.append(" ".join(cur).strip())
#                     cur, cur_len = [], 0
#                 cur.append(s); cur_len += len(s) + 1
#             if cur:
#                 segs.append(" ".join(cur).strip())
#     return segs

# def run_coding(file_name: str, content: bytes, max_chars: int) -> str:
#     text = read_docx_bytes(content) if file_name.lower().endswith(".docx") else read_txt_bytes(content)
#     segments = segment_for_coding(text, max_chars=max_chars)
#     lines = [f"=== KII CODED REPORT â€” {file_name} ===", ""]
#     progress = st.progress(0, text="Coding segmentsâ€¦")
#     for i, seg in enumerate(segments, start=1):
#         out = call_azure(build_kII_prompt(seg, KII_TAXONOMY))
#         lines.append(f"[Segment {i}]"); lines.append(out); lines.append("")
#         progress.progress(i / len(segments), text=f"Coded {i}/{len(segments)}")
#         time.sleep(0.05)
#     return "\n".join(lines)

# # â”€â”€â”€â”€â”€â”€â”€â”€â”€ Tabs UI (NO sidebar) â”€â”€â”€â”€â”€â”€â”€â”€â”€
# tab1, tab2 = st.tabs(["ğŸ§  Insights Extraction", "ğŸ·ï¸ Codebook Coding (KII)"])

# with tab1:
#     st.subheader("Transcript â†’ Insights Extraction")
#     uploads = st.file_uploader("Upload transcript files (.docx / .txt)", type=["docx","txt"], accept_multiple_files=True, key="ins_uploads")
#     go = st.button("ğŸš€ Generate Insights", key="ins_go")
#     if go:
#         if not uploads:
#             st.warning("Please upload at least one file.")
#         else:
#             for up in uploads:
#                 st.write(f"**Processing:** `{up.name}`")
#                 report = run_insights(up.name, up.read())
#                 with st.expander(f"ğŸ“„ Insights: {up.name}", expanded=True):
#                     st.text(report)
#                 st.download_button(
#                     label=f"ğŸ’¾ Download â€” {os.path.splitext(up.name)[0]}_INSIGHTS.txt",
#                     data=report.encode("utf-8"),
#                     file_name=f"{os.path.splitext(up.name)[0]}_INSIGHTS.txt",
#                     mime="text/plain"
#                 )

# with tab2:
#     st.subheader("KII Transcript â†’ Thematic Coding (Plain Text)")
#     seg_size = st.slider("Max segment size (chars)", 800, 3000, 1800, 100, key="code_seg")
#     uploads2 = st.file_uploader("Upload transcript files (.docx / .txt)", type=["docx","txt"], accept_multiple_files=True, key="code_uploads")
#     go2 = st.button("ğŸš€ Run Coding", key="code_go")
#     if go2:
#         if not uploads2:
#             st.warning("Please upload at least one file.")
#         else:
#             for up in uploads2:
#                 st.write(f"**Processing:** `{up.name}`")
#                 report = run_coding(up.name, up.read(), max_chars=seg_size)
#                 with st.expander(f"ğŸ“„ Coded Report: {up.name}", expanded=True):
#                     st.text(report)
#                 st.download_button(
#                     label=f"ğŸ’¾ Download â€” {os.path.splitext(up.name)[0]}_KII_CODED.txt",
#                     data=report.encode("utf-8"),
#                     file_name=f"{os.path.splitext(up.name)[0]}_KII_CODED.txt",
#                     mime="text/plain"
#                 )
# --------------------------------------------------------------------------------------------------------------------------------------------


# # app.py
# import os, io, re, time, string
# from typing import List, Dict
# import streamlit as st
# from docx import Document
# from openai import AzureOpenAI
# from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type

# # â”€â”€â”€â”€â”€â”€â”€â”€â”€ Page â”€â”€â”€â”€â”€â”€â”€â”€â”€
# st.set_page_config(page_title="NF Transcript Toolkit", layout="wide")
# st.title("ğŸŒ¾ NF Transcript Toolkit")
# st.caption("Use the tabs below to switch between **Insights Extraction** and **Codebook Coding (KII)**. Azure keys are read from Streamlit Secrets.")

# # â”€â”€â”€â”€â”€â”€â”€â”€â”€ Azure Client (from Secrets) â”€â”€â”€â”€â”€â”€â”€â”€â”€
# def init_client():
#     try:
#         client = AzureOpenAI(
#             api_key=st.secrets["AZURE_OPENAI_API_KEY"],
#             api_version=st.secrets.get("AZURE_OPENAI_API_VERSION", "2024-10-21"),
#             azure_endpoint=st.secrets["AZURE_OPENAI_ENDPOINT"].strip()  # e.g. https://<resource>.openai.azure.com/
#         )
#         deployment = st.secrets.get("DEPLOYMENT", "gpt-4o")
#         return client, deployment
#     except Exception:
#         st.error("Azure OpenAI secrets missing/invalid. Set them in App â†’ Settings â†’ Secrets.")
#         st.stop()

# client, DEPLOYMENT = init_client()

# @retry(wait=wait_exponential(multiplier=1, min=1, max=20),
#        stop=stop_after_attempt(6),
#        retry=retry_if_exception_type(Exception))
# def call_azure(prompt: str) -> str:
#     res = client.chat.completions.create(
#         model=DEPLOYMENT,
#         messages=[{"role": "user", "content": prompt}],
#         temperature=0
#     )
#     return res.choices[0].message.content.strip()

# # â”€â”€â”€â”€â”€â”€â”€â”€â”€ Shared Helpers â”€â”€â”€â”€â”€â”€â”€â”€â”€
# def read_docx_bytes(b: bytes) -> str:
#     doc = Document(io.BytesIO(b))
#     return "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())

# def read_txt_bytes(b: bytes) -> str:
#     return b.decode("utf-8", errors="ignore")

# def clean_text(s: str) -> str:
#     return ''.join(c for c in s if c in string.printable)

# def chunk_text(text: str, max_chars: int = 6000, overlap: int = 400) -> List[str]:
#     text = re.sub(r'\r\n?', '\n', text)
#     if len(text) <= max_chars:
#         return [text]
#     chunks, i = [], 0
#     while i < len(text):
#         chunk = text[i:i+max_chars]
#         lb = chunk.rfind("\n\n")
#         if lb > max_chars * 0.6:
#             chunk = chunk[:lb]; i += lb
#         else:
#             i += max_chars
#         chunks.append(chunk.strip())
#         i -= overlap
#         if i < 0: i = 0
#     # de-dup by first 200 chars
#     seen, dedup = set(), []
#     for c in chunks:
#         k = c[:200]
#         if k not in seen:
#             seen.add(k); dedup.append(c)
#     return dedup

# def normalize_bullets(lines: List[str]) -> List[str]:
#     out, seen = [], set()
#     for ln in lines:
#         s = re.sub(r'\s+', ' ', ln.strip().rstrip('.'))
#         if not s: continue
#         key = s.lower()
#         if key not in seen:
#             seen.add(key)
#             out.append(s if s.endswith('.') else s + '.')
#     return out

# # â”€â”€â”€â”€â”€â”€â”€â”€â”€ Mode 1: Insights Extraction â”€â”€â”€â”€â”€â”€â”€â”€â”€
# def chunk_prompt(chunk: str) -> str:
#     return f"""
# You are a senior qualitative researcher. Read the TEXT and extract clear, non-generic insights.
# Focus on perceptions, experiences, outcomes, barriers, enablers, and actionable suggestions.
# Return ONLY the following two sections as plain text (no JSON, no Markdown):

# INSIGHTS:
# - <8â€“15 concise bullets, each a standalone insight; avoid repetition; specific to this text>

# QUOTES:
# - <5â€“10 short verbatim quotes (<=25 words each); keep any speaker tags like 'P3:' if present>

# TEXT:
# \"\"\"{chunk}\"\"\"
# """

# def synthesis_prompt(all_insights: List[str], all_quotes: List[str], fname: str) -> str:
#     ins = "\n".join(f"- {i}" for i in all_insights)
#     qts = "\n".join(f"- {q}" for q in all_quotes)
#     return f"""
# You will merge bullets from multiple chunks into a single clean report for {fname}.
# Deduplicate, group logically, and keep only the strongest items. Be specific; no filler.
# Return ONLY plain text (no JSON/Markdown), exactly in this template:

# INSIGHTS (Top 10â€“20):
# - <bullet>
# - <bullet>
# ...

# QUOTES (5â€“12 strongest):
# - <short quote>
# - <short quote>
# ...

# BULLET POOL:
# {ins}

# QUOTE POOL:
# {qts}
# """

# def run_insights(file_name: str, content: bytes) -> str:
#     text = read_docx_bytes(content) if file_name.lower().endswith(".docx") else read_txt_bytes(content)
#     text = clean_text(text)
#     chunks = chunk_text(text, max_chars=6000, overlap=400)
#     pooled_insights, pooled_quotes = [], []

#     progress = st.progress(0, text="Extracting segmentsâ€¦")
#     for i, ch in enumerate(chunks, start=1):
#         resp = call_azure(chunk_prompt(ch))
#         parts = re.split(r'\bQUOTES:\s*', resp, flags=re.I)
#         insights_block = parts[0]
#         quotes_block = parts[1] if len(parts) > 1 else ""
#         ins = re.findall(r'^\s*-\s+(.*)', insights_block, flags=re.M)
#         qts = re.findall(r'^\s*-\s+(.*)', quotes_block, flags=re.M)
#         pooled_insights.extend(ins); pooled_quotes.extend(qts)
#         progress.progress(i / len(chunks), text=f"Extracted {i}/{len(chunks)}")

#     pooled_insights = normalize_bullets(pooled_insights)
#     pooled_quotes = normalize_bullets(pooled_quotes)
#     final_text = call_azure(synthesis_prompt(pooled_insights, pooled_quotes, file_name)).strip()
#     return f"INSIGHTS REPORT â€” {file_name}\n\n{final_text}"

# # â”€â”€â”€â”€â”€â”€â”€â”€â”€ Mode 2: Codebook Coding (KII) â”€â”€â”€â”€â”€â”€â”€â”€â”€
# KII_TAXONOMY: Dict[str, Dict[str, str]] = {
#     "Respondent Background (BCK)": {
#         "BCK_Role": "Role of respondent (farmer/CRP/leader).",
#         "BCK_Experience": "Years of farming / NF experience."
#     },
#     "Introduction and Spread of NF (SPR)": {
#         "SPR_Introduced_By": "Who introduced NF (CSA/govt/others).",
#         "SPR_Response": "Communityâ€™s initial reactions.",
#         "SPR_EarlyAdopters": "Presence/influence of early adopters."
#     },
#     "Adoption Patterns (ADO)": {
#         "ADO_Numbers": "Adopters initially vs now.",
#         "ADO_Types": "Types of farmers adopting (small/marginal/large).",
#         "ADO_DropoutReasons": "Reasons for dropping out."
#     },
#     "NF Practices and Inputs (PRC)": {
#         "PRC_BiologicalInputs": "Inputs like FYM/neem/panchagavya.",
#         "PRC_InputAccess": "Access/availability of inputs.",
#         "PRC_CropTypes": "Crops under NF / cropping changes."
#     },
#     "Challenges and Barriers (CHL)": {
#         "CHL_Labor": "Labor-related constraints.",
#         "CHL_Yield": "Yield-related concerns.",
#         "CHL_Market": "Market/premium issues.",
#         "CHL_Social": "Social stigma/peer pressure.",
#         "CHL_HealthFoodChange": "Diet/health change issues."
#     },
#     "Support Systems (SUP)": {
#         "SUP_CSA": "CSA training/demos/CRP support.",
#         "SUP_GovtSupport": "Govt schemes/benefits/gaps.",
#         "SUP_Suggestions": "Suggestions to improve support."
#     },
#     "Benefits of NF (BEN)": {
#         "BEN_CostReduction": "Reduced input costs.",
#         "BEN_SoilHealth": "Soil improvements.",
#         "BEN_HumanHealth": "Human health improvements.",
#         "BEN_IncomeStability": "Income stability/parity."
#     },
#     "Ecosystem and Policy (ECO)": {
#         "ECO_MarketPolicy": "Price support / exclusive markets.",
#         "ECO_EcosystemChange": "System/political leadership asks.",
#         "ECO_InstitutionalLink": "Links to govt institutions."
#     },
#     "Future Direction and Continuity (FUT)": {
#         "FUT_ExposureVisits": "Exposure/knowledge sharing.",
#         "FUT_TrainingNeed": "Further/refresher training needs.",
#         "FUT_Infrastructure": "Storage/input centres/infrastructure."
#     }
# }

# def taxonomy_lines(tax: Dict[str, Dict[str,str]]) -> str:
#     return "\n".join(f"{theme} â†’ {', '.join(subs.keys())}" for theme, subs in tax.items())

# def build_kII_prompt(segment: str, tax: Dict[str,Dict[str,str]]) -> str:
#     tax_text = taxonomy_lines(tax)
#     return f"""
# You are a qualitative coding assistant. Classify the KII excerpt into the taxonomy BELOW.
# Pick exactly ONE theme and one or more subcodes from THAT theme only.
# Then write a short insight and one short verbatim quote from the text (<=25 words).

# TAXONOMY:
# {tax_text}

# TEXT:
# \"\"\"{segment}\"\"\"

# Return PLAIN TEXT only (no JSON, no backticks), exactly in this format:

# Theme: <one theme name exactly as in taxonomy>
# Subcodes: <comma-separated subcodes from the chosen theme>
# Insight: <one concise, specific sentence>
# Quote: "<short verbatim quote from TEXT, <=25 words>"

# (Do not add any extra lines or sections.)
# """.strip()

# def segment_for_coding(text: str, max_chars: int = 1800) -> List[str]:
#     text = re.sub(r'\r\n?', '\n', text)
#     blocks = [b.strip() for b in re.split(r'\n{2,}', text) if b.strip()]
#     segs: List[str] = []
#     for blk in blocks:
#         if len(blk) <= max_chars:
#             segs.append(blk)
#         else:
#             sentences = re.split(r'(?<=[.!?])\s+', blk)
#             cur, cur_len = [], 0
#             for s in sentences:
#                 if cur_len + len(s) > max_chars and cur:
#                     segs.append(" ".join(cur).strip()); cur, cur_len = [], 0
#                 cur.append(s); cur_len += len(s) + 1
#             if cur:
#                 segs.append(" ".join(cur).strip())
#     return segs

# def run_coding(file_name: str, content: bytes, max_chars: int) -> str:
#     text = read_docx_bytes(content) if file_name.lower().endswith(".docx") else read_txt_bytes(content)
#     segments = segment_for_coding(text, max_chars=max_chars)
#     lines = [f"=== KII CODED REPORT â€” {file_name} ===", ""]
#     progress = st.progress(0, text="Coding segmentsâ€¦")
#     for i, seg in enumerate(segments, start=1):
#         out = call_azure(build_kII_prompt(seg, KII_TAXONOMY))
#         lines.append(f"[Segment {i}]"); lines.append(out); lines.append("")
#         progress.progress(i / len(segments), text=f"Coded {i}/{len(segments)}")
#         time.sleep(0.05)
#     return "\n".join(lines)

# # â”€â”€â”€â”€â”€â”€â”€â”€â”€ Tabs UI (unique keys) â”€â”€â”€â”€â”€â”€â”€â”€â”€
# tab1, tab2 = st.tabs(["ğŸ§  Insights Extraction", "ğŸ·ï¸ Codebook Coding (KII)"])

# with tab1:
#     st.subheader("Transcript â†’ Insights Extraction")
#     uploads = st.file_uploader(
#         "Upload transcript files (.docx / .txt)",
#         type=["docx","txt"],
#         accept_multiple_files=True,
#         key="ins_uploads"
#     )
#     go_insights = st.button("ğŸš€ Generate Insights", key="btn_insights")
#     if go_insights:
#         if not uploads:
#             st.warning("Please upload at least one file.")
#         else:
#             for up in uploads:
#                 st.write(f"**Processing:** `{up.name}`")
#                 report = run_insights(up.name, up.read())
#                 with st.expander(f"ğŸ“„ Insights: {up.name}", expanded=True):
#                     st.text(report)
#                 st.download_button(
#                     label=f"ğŸ’¾ Download â€” {os.path.splitext(up.name)[0]}_INSIGHTS.txt",
#                     data=report.encode("utf-8"),
#                     file_name=f"{os.path.splitext(up.name)[0]}_INSIGHTS.txt",
#                     mime="text/plain",
#                     key=f"dl_ins_{up.name}"
#                 )

# with tab2:
#     st.subheader("KII Transcript â†’ Thematic Coding (Plain Text)")
#     seg_size = st.slider("Max segment size (chars)", 800, 3000, 1800, 100, key="code_seg")
#     uploads2 = st.file_uploader(
#         "Upload transcript files (.docx / .txt)",
#         type=["docx","txt"],
#         accept_multiple_files=True,
#         key="code_uploads"
#     )
#     go_coding = st.button("ğŸš€ Run Coding", key="btn_coding")
#     if go_coding:
#         if not uploads2:
#             st.warning("Please upload at least one file.")
#         else:
#             for up in uploads2:
#                 st.write(f"**Processing:** `{up.name}`")
#                 report = run_coding(up.name, up.read(), max_chars=seg_size)
#                 with st.expander(f"ğŸ“„ Coded Report: {up.name}", expanded=True):
#                     st.text(report)
#                 st.download_button(
#                     label=f"ğŸ’¾ Download â€” {os.path.splitext(up.name)[0]}_KII_CODED.txt",
#                     data=report.encode("utf-8"),
#                     file_name=f"{os.path.splitext(up.name)[0]}_KII_CODED.txt",
#                     mime="text/plain",
#                     key=f"dl_code_{up.name}"
#                 )


# app.py
import os, io, re, time, string
from typing import List, Dict
import streamlit as st
from docx import Document
from openai import AzureOpenAI
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type

# â”€â”€â”€â”€â”€â”€â”€â”€â”€ Page â”€â”€â”€â”€â”€â”€â”€â”€â”€
st.set_page_config(page_title="NF Transcript Toolkit", layout="wide")
st.title("ğŸŒ¾ NF Transcript Toolkit")
st.caption("Use the tabs below to switch between **Insights Extraction** and **Codebook Coding (KII)**. Azure keys are read from Streamlit Secrets.")

# â”€â”€â”€â”€â”€â”€â”€â”€â”€ Session State (persist results across reruns & tab switches) â”€â”€â”€â”€â”€â”€â”€â”€â”€
if "insights_reports" not in st.session_state:
    # { filename: {"report": str, "timestamp": float} }
    st.session_state.insights_reports = {}
if "coding_reports" not in st.session_state:
    # { filename: {"report": str, "timestamp": float, "seg_size": int} }
    st.session_state.coding_reports = {}

# â”€â”€â”€â”€â”€â”€â”€â”€â”€ Azure Client (from Secrets) â”€â”€â”€â”€â”€â”€â”€â”€â”€
def init_client():
    try:
        client = AzureOpenAI(
            api_key=st.secrets["AZURE_OPENAI_API_KEY"],
            api_version=st.secrets.get("AZURE_OPENAI_API_VERSION", "2024-10-21"),
            azure_endpoint=st.secrets["AZURE_OPENAI_ENDPOINT"].strip()  # e.g. https://<resource>.openai.azure.com/
        )
        deployment = st.secrets.get("DEPLOYMENT", "gpt-4o")
        return client, deployment
    except Exception:
        st.error("Azure OpenAI secrets missing/invalid. Set them in App â†’ Settings â†’ Secrets.")
        st.stop()

client, DEPLOYMENT = init_client()

@retry(wait=wait_exponential(multiplier=1, min=1, max=20),
       stop=stop_after_attempt(6),
       retry=retry_if_exception_type(Exception))
def call_azure(prompt: str) -> str:
    res = client.chat.completions.create(
        model=DEPLOYMENT,
        messages=[{"role": "user", "content": prompt}],
        temperature=0
    )
    return res.choices[0].message.content.strip()

# â”€â”€â”€â”€â”€â”€â”€â”€â”€ Shared Helpers â”€â”€â”€â”€â”€â”€â”€â”€â”€
def read_docx_bytes(b: bytes) -> str:
    doc = Document(io.BytesIO(b))
    return "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())

def read_txt_bytes(b: bytes) -> str:
    return b.decode("utf-8", errors="ignore")

def clean_text(s: str) -> str:
    return ''.join(c for c in s if c in string.printable)

def chunk_text(text: str, max_chars: int = 6000, overlap: int = 400) -> List[str]:
    text = re.sub(r'\r\n?', '\n', text)
    if len(text) <= max_chars:
        return [text]
    chunks, i = [], 0
    while i < len(text):
        chunk = text[i:i+max_chars]
        lb = chunk.rfind("\n\n")
        if lb > max_chars * 0.6:
            chunk = chunk[:lb]; i += lb
        else:
            i += max_chars
        chunks.append(chunk.strip())
        i -= overlap
        if i < 0: i = 0
    # de-dup by first 200 chars
    seen, dedup = set(), []
    for c in chunks:
        k = c[:200]
        if k not in seen:
            seen.add(k); dedup.append(c)
    return dedup

def normalize_bullets(lines: List[str]) -> List[str]:
    out, seen = [], set()
    for ln in lines:
        s = re.sub(r'\s+', ' ', ln.strip().rstrip('.'))
        if not s: continue
        key = s.lower()
        if key not in seen:
            seen.add(key)
            out.append(s if s.endswith('.') else s + '.')
    return out

# â”€â”€â”€â”€â”€â”€â”€â”€â”€ Mode 1: Insights Extraction â”€â”€â”€â”€â”€â”€â”€â”€â”€
def chunk_prompt(chunk: str) -> str:
    return f"""
You are a senior qualitative researcher. Read the TEXT and extract clear, non-generic insights.
Focus on perceptions, experiences, outcomes, barriers, enablers, and actionable suggestions.
Return ONLY the following two sections as plain text (no JSON, no Markdown):

INSIGHTS:
- <8â€“15 concise bullets, each a standalone insight; avoid repetition; specific to this text>

QUOTES:
- <5â€“10 short verbatim quotes (<=25 words each); keep any speaker tags like 'P3:' if present>

TEXT:
\"\"\"{chunk}\"\"\"
"""

def synthesis_prompt(all_insights: List[str], all_quotes: List[str], fname: str) -> str:
    ins = "\n".join(f"- {i}" for i in all_insights)
    qts = "\n".join(f"- {q}" for q in all_quotes)
    return f"""
You will merge bullets from multiple chunks into a single clean report for {fname}.
Deduplicate, group logically, and keep only the strongest items. Be specific; no filler.
Return ONLY plain text (no JSON/Markdown), exactly in this template:

INSIGHTS (Top 10â€“20):
- <bullet>
- <bullet>
...

QUOTES (5â€“12 strongest):
- <short quote>
- <short quote>
...

BULLET POOL:
{ins}

QUOTE POOL:
{qts}
"""

def run_insights(file_name: str, content: bytes) -> str:
    text = read_docx_bytes(content) if file_name.lower().endswith(".docx") else read_txt_bytes(content)
    text = clean_text(text)
    chunks = chunk_text(text, max_chars=6000, overlap=400)
    pooled_insights, pooled_quotes = [], []

    progress = st.progress(0, text="Extracting segmentsâ€¦")
    for i, ch in enumerate(chunks, start=1):
        resp = call_azure(chunk_prompt(ch))
        parts = re.split(r'\bQUOTES:\s*', resp, flags=re.I)
        insights_block = parts[0]
        quotes_block = parts[1] if len(parts) > 1 else ""
        ins = re.findall(r'^\s*-\s+(.*)', insights_block, flags=re.M)
        qts = re.findall(r'^\s*-\s+(.*)', quotes_block, flags=re.M)
        pooled_insights.extend(ins); pooled_quotes.extend(qts)
        progress.progress(i / len(chunks), text=f"Extracted {i}/{len(chunks)}")

    pooled_insights = normalize_bullets(pooled_insights)
    pooled_quotes = normalize_bullets(pooled_quotes)
    final_text = call_azure(synthesis_prompt(pooled_insights, pooled_quotes, file_name)).strip()
    return f"INSIGHTS REPORT â€” {file_name}\n\n{final_text}"

# â”€â”€â”€â”€â”€â”€â”€â”€â”€ Mode 2: Codebook Coding (KII) â”€â”€â”€â”€â”€â”€â”€â”€â”€
KII_TAXONOMY: Dict[str, Dict[str, str]] = {
    "Respondent Background (BCK)": {"BCK_Role": "...", "BCK_Experience": "..."},
    "Introduction and Spread of NF (SPR)": {"SPR_Introduced_By": "...", "SPR_Response": "...", "SPR_EarlyAdopters": "..."},
    "Adoption Patterns (ADO)": {"ADO_Numbers": "...", "ADO_Types": "...", "ADO_DropoutReasons": "..."},
    "NF Practices and Inputs (PRC)": {"PRC_BiologicalInputs": "...", "PRC_InputAccess": "...", "PRC_CropTypes": "..."},
    "Challenges and Barriers (CHL)": {"CHL_Labor": "...", "CHL_Yield": "...", "CHL_Market": "...", "CHL_Social": "...", "CHL_HealthFoodChange": "..."},
    "Support Systems (SUP)": {"SUP_CSA": "...", "SUP_GovtSupport": "...", "SUP_Suggestions": "..."},
    "Benefits of NF (BEN)": {"BEN_CostReduction": "...", "BEN_SoilHealth": "...", "BEN_HumanHealth": "...", "BEN_IncomeStability": "..."},
    "Ecosystem and Policy (ECO)": {"ECO_MarketPolicy": "...", "ECO_EcosystemChange": "...", "ECO_InstitutionalLink": "..."},
    "Future Direction and Continuity (FUT)": {"FUT_ExposureVisits": "...", "FUT_TrainingNeed": "...", "FUT_Infrastructure": "..."},
}

def taxonomy_lines(tax: Dict[str, Dict[str,str]]) -> str:
    return "\n".join(f"{theme} â†’ {', '.join(subs.keys())}" for theme, subs in tax.items())

def build_kII_prompt(segment: str, tax: Dict[str,Dict[str,str]]) -> str:
    tax_text = taxonomy_lines(tax)
    return f"""
You are a qualitative coding assistant. Classify the KII excerpt into the taxonomy BELOW.
Pick exactly ONE theme and one or more subcodes from THAT theme only.
Then write a short insight and one short verbatim quote from the text (<=25 words).

TAXONOMY:
{tax_text}

TEXT:
\"\"\"{segment}\"\"\"

Return PLAIN TEXT only (no JSON, no backticks), exactly in this format:

Theme: <one theme name exactly as in taxonomy>
Subcodes: <comma-separated subcodes from the chosen theme>
Insight: <one concise, specific sentence>
Quote: "<short verbatim quote from TEXT, <=25 words>"

(Do not add any extra lines or sections.)
""".strip()

def segment_for_coding(text: str, max_chars: int = 1800) -> List[str]:
    text = re.sub(r'\r\n?', '\n', text)
    blocks = [b.strip() for b in re.split(r'\n{2,}', text) if b.strip()]
    segs: List[str] = []
    for blk in blocks:
        if len(blk) <= max_chars:
            segs.append(blk)
        else:
            sentences = re.split(r'(?<=[.!?])\s+', blk)
            cur, cur_len = [], 0
            for s in sentences:
                if cur_len + len(s) > max_chars and cur:
                    segs.append(" ".join(cur).strip()); cur, cur_len = [], 0
                cur.append(s); cur_len += len(s) + 1
            if cur:
                segs.append(" ".join(cur).strip())
    return segs

def run_coding(file_name: str, content: bytes, max_chars: int) -> str:
    text = read_docx_bytes(content) if file_name.lower().endswith(".docx") else read_txt_bytes(content)
    segments = segment_for_coding(text, max_chars=max_chars)
    lines = [f"=== KII CODED REPORT â€” {file_name} ===", ""]
    progress = st.progress(0, text="Coding segmentsâ€¦")
    for i, seg in enumerate(segments, start=1):
        out = call_azure(build_kII_prompt(seg, KII_TAXONOMY))
        lines.append(f"[Segment {i}]"); lines.append(out); lines.append("")
        progress.progress(i / len(segments), text=f"Coded {i}/{len(segments)}")
        time.sleep(0.05)
    return "\n".join(lines)

# â”€â”€â”€â”€â”€â”€â”€â”€â”€ Tabs UI (persisted results) â”€â”€â”€â”€â”€â”€â”€â”€â”€
tab1, tab2 = st.tabs(["ğŸ§  Insights Extraction", "ğŸ·ï¸ Codebook Coding (KII)"])

with tab1:
    st.subheader("Transcript â†’ Insights Extraction")
    uploads = st.file_uploader(
        "Upload transcript files (.docx / .txt)",
        type=["docx","txt"],
        accept_multiple_files=True,
        key="ins_uploads"
    )

    colA, colB = st.columns([1,1])
    with colA:
        go_insights = st.button("ğŸš€ Generate Insights", key="btn_insights")
    with colB:
        if st.button("ğŸ§¹ Clear Results", key="btn_clear_ins"):
            st.session_state.insights_reports.clear()
            st.success("Cleared insights results.")

    # Generate & SAVE to session_state
    if go_insights:
        if not uploads:
            st.warning("Please upload at least one file.")
        else:
            for up in uploads:
                st.write(f"**Processing:** `{up.name}`")
                report = run_insights(up.name, up.read())
                st.session_state.insights_reports[up.name] = {
                    "report": report,
                    "timestamp": time.time()
                }

    # Always DISPLAY anything already saved
    if st.session_state.insights_reports:
        st.markdown("### Saved Insights Reports")
        # show newest first
        for fname, payload in sorted(st.session_state.insights_reports.items(),
                                     key=lambda kv: kv[1]["timestamp"],
                                     reverse=True):
            with st.expander(f"ğŸ“„ Insights: {fname}", expanded=False):
                st.text(payload["report"])
            st.download_button(
                label=f"ğŸ’¾ Download â€” {os.path.splitext(fname)[0]}_INSIGHTS.txt",
                data=payload["report"].encode("utf-8"),
                file_name=f"{os.path.splitext(fname)[0]}_INSIGHTS.txt",
                mime="text/plain",
                key=f"dl_ins_{fname}"
            )
    else:
        st.info("No insights yet. Upload files and click **Generate Insights**.")

with tab2:
    st.subheader("KII Transcript â†’ Thematic Coding (Plain Text)")
    seg_size = st.slider("Max segment size (chars)", 800, 3000, 1800, 100, key="code_seg")
    uploads2 = st.file_uploader(
        "Upload transcript files (.docx / .txt)",
        type=["docx","txt"],
        accept_multiple_files=True,
        key="code_uploads"
    )

    colC, colD = st.columns([1,1])
    with colC:
        go_coding = st.button("ğŸš€ Run Coding", key="btn_coding")
    with colD:
        if st.button("ğŸ§¹ Clear Results", key="btn_clear_code"):
            st.session_state.coding_reports.clear()
            st.success("Cleared coding results.")

    # Generate & SAVE to session_state
    if go_coding:
        if not uploads2:
            st.warning("Please upload at least one file.")
        else:
            for up in uploads2:
                st.write(f"**Processing:** `{up.name}`")
                report = run_coding(up.name, up.read(), max_chars=seg_size)
                st.session_state.coding_reports[up.name] = {
                    "report": report,
                    "seg_size": seg_size,
                    "timestamp": time.time()
                }

    # Always DISPLAY anything already saved
    if st.session_state.coding_reports:
        st.markdown("### Saved Coded Reports")
        for fname, payload in sorted(st.session_state.coding_reports.items(),
                                     key=lambda kv: kv[1]["timestamp"],
                                     reverse=True):
            with st.expander(f"ğŸ·ï¸ Coded Report: {fname}", expanded=False):
                st.text(payload["report"])
            st.download_button(
                label=f"ğŸ’¾ Download â€” {os.path.splitext(fname)[0]}_KII_CODED.txt",
                data=payload["report"].encode("utf-8"),
                file_name=f"{os.path.splitext(fname)[0]}_KII_CODED.txt",
                mime="text/plain",
                key=f"dl_code_{fname}"
            )
    else:
        st.info("No coded outputs yet. Upload files and click **Run Coding**.")

